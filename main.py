from flask import Flask,request,render_template,send_file
from Preprocessing import validation
from Stocks_scraper import StockInfo
from Directory_opps import dir_opps
from io import BytesIO
import docx
from DB import Database_Operations
from Content import Message
from Reports import PriorTOdoc
import os

app = Flask(__name__)

@app.route("/",methods=['GET','POST'])
def index():
    return render_template('index.html')

@app.route('/values',methods=['POST'])

def intution(directory = 'C:\\Users\\sahas\\PycharmProjects\\Stocks\\'):
    Company_name = request.form['company_name']
    Stock_Accronym = request.form['Stock_name']
    start_year = request.form['start_date'].split('-')
    end_year = request.form['end_date'].split('-')
    # Scraping the stock data by using Stock info method of stock scraper class
    stock=StockInfo(start_year, end_year)
    stock.Stock_data(Stock_Accronym)
    dir_=dir_opps(Stock_Accronym+'.csv', Company_name)
    dir_.rm_dir()
    # Validating the dataset (i.e) preprocess
    v=validation(Stock_Accronym, Company_name)
    v.DataPreprocessing()
    # Generating Reports
    document = docx.Document()
    f = BytesIO()
    # Add Heading1
    document.add_heading(Company_name.capitalize())
    # Add Paragraph
    document.add_paragraph(Message().Introduction().format(StockInfo(start_year,end_year).Date()[1][0]-StockInfo(start_year,end_year).Date()[0][0],StockInfo(start_year,end_year).Date()[0][0],StockInfo(start_year,end_year).Date()[1][0]))
    # add heading2
    document.add_heading(Message().Visuals())
    # Add picture visuals
    for pic_ in PriorTOdoc(Stock_Accronym+'_validated_.csv',Company_name).Visuvalization():
        document.add_picture(pic_)
        os.remove(directory+pic_)
    # Add Heading 3
    document.add_heading('Spread')
    # Add paragraph
    document.add_paragraph(Message().Spread())
    # Add picture for spread
    for pic_ in PriorTOdoc(Stock_Accronym+'_validated_.csv',Company_name).Spread():
        document.add_picture(pic_)
        os.remove(directory+pic_)
    # Add Heading
    document.add_heading('Volatility')
    # Add Paragraph
    document.add_paragraph(Message().volatility())
    document.add_paragraph(PriorTOdoc(Stock_Accronym+'_validated_.csv',Company_name).volatility()[0])
    # Add Bar chart
    for pic_ in PriorTOdoc(Stock_Accronym + '_validated_.csv', Company_name).volatility_graph():
        document.add_picture(pic_)
        os.remove(directory + pic_)

    # Add Heading
    document.add_heading('Daily Returns')
    # Add Paragraph
    document.add_paragraph(Message().returns())
    # Add picture
    for pic_ in PriorTOdoc(Stock_Accronym+'_validated_.csv',Company_name).Returns()[1]:
        document.add_picture(pic_)
        os.remove(directory+pic_)

    # Add sub heading
    document.add_heading('Daily volatility for returns')
    # Add paragraph
    document.add_paragraph(PriorTOdoc(Stock_Accronym+'_validated_.csv',Company_name).dailyReturnvolatility())
    # Add Bar chart
    for pic_ in PriorTOdoc(Stock_Accronym + '_validated_.csv', Company_name).returns_graph():
        document.add_picture(pic_)
        os.remove(directory + pic_)
    # Add Heading
    document.add_heading('Moving Average')
    # Add paragraph
    document.add_paragraph(Message().moving_average())
    # Add picture
    for pic_ in PriorTOdoc(Stock_Accronym+'_validated_.csv',Company_name).moving_average():
        document.add_picture(pic_)
        os.remove(directory+pic_)
    # Add heading
    document.add_heading('Correlation Plots')
    # Add paragraph
    document.add_paragraph(Message().auto_correlation())
    # Add picture auto_corr
    for pic_ in PriorTOdoc(Stock_Accronym + '_validated_.csv', Company_name).correlation_plots():
        document.add_picture(pic_)
        os.remove(directory+pic_)
    # Add heading
    document.add_heading('Predictions - LSTM')
    # Add paragraph
    document.add_paragraph(Message().lstm())
    # Add picture
    file_name=PriorTOdoc(Stock_Accronym +'_validated_.csv', Company_name).LSTM()[0]
    document.add_picture(file_name)
    os.remove(directory+file_name)
    document.save(f)
    f.seek(0)

    # Dropping the Collection from the Database
    Database_Operations(database_name='financialdata', Table_name=Stock_Accronym.split('.')[0], file_name=Stock_Accronym+'.csv').drop_databases(database_name='financialdata')
    return send_file(f,as_attachment=True,attachment_filename=Company_name+'.docx')


app.run(port=8050,debug=True)
